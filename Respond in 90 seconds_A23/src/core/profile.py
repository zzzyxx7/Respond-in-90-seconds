"""
Profile 生成器 — 从模板文件或自然语言描述生成提取配置
"""

import json
import logging
import os
import re
from pathlib import Path
from typing import Optional, Dict, Any

from src.core.template_detector import detect_template_structure
from src.core.alias import resolve_field_names
from src.core.instruction_filters import parse_date_range_from_instruction

logger = logging.getLogger(__name__)


# ── 字段类型自动推断 ────────────────────────────────────────────────────────

# 从 field_normalization_rules.json 加载已知字段→类型映射（启动时一次性加载）
_RULES_TYPE_CACHE: Optional[Dict[str, str]] = None


def _load_rules_type_map() -> Dict[str, str]:
    """从 field_normalization_rules.json 读取 fields 段的字段→类型映射"""
    global _RULES_TYPE_CACHE
    if _RULES_TYPE_CACHE is not None:
        return _RULES_TYPE_CACHE
    _RULES_TYPE_CACHE = {}
    try:
        rules_path = Path(__file__).parent.parent / "knowledge" / "field_normalization_rules.json"
        with open(rules_path, "r", encoding="utf-8") as f:
            rules = json.load(f)
        for field_name, field_rule in rules.get("fields", {}).items():
            if isinstance(field_rule, dict) and "type" in field_rule:
                _RULES_TYPE_CACHE[field_name] = field_rule["type"]
    except Exception:
        pass
    return _RULES_TYPE_CACHE


def _infer_field_type(field_name: str, unit: str = "") -> str:
    """根据字段名和单位自动推断字段类型

    推断策略（按优先级）：
    1. 有单位 → 通过单位内容判断（带数量/金额单位=number，带%=percentage）
    2. 字段名在 field_normalization_rules.json 中有定义 → 用定义的类型
    3. 字段名通过别名解析后在规则中有定义 → 继承类型
    4. 兜底 → text
    """
    # 1. 单位推断：有单位的字段几乎一定是数值或百分比
    if unit:
        clean_unit = unit.strip()
        if "%" in clean_unit or "％" in clean_unit or "‰" in clean_unit:
            return "percentage"
        # 有任何非空单位 → 数值类型（通用规则，覆盖元/万/吨/km等所有场景）
        if clean_unit:
            return "number"

    # 2. 查已知字段规则
    rules_map = _load_rules_type_map()
    if field_name in rules_map:
        return rules_map[field_name]

    # 3. 通过别名解析后再查规则
    try:
        from src.core.alias import load_alias_map, resolve_field_name
        alias_map = load_alias_map()
        canonical = resolve_field_name(field_name, alias_map, fuzzy_threshold=80)
        if canonical != field_name and canonical in rules_map:
            return rules_map[canonical]
    except Exception:
        pass

    return "text"


def _enrich_fields(raw_fields: list, resolved_fields: list) -> tuple:
    """批量构建字段定义并自动推断关键字段

    Returns:
        (fields_list, dedup_key_fields)
        关键字段识别策略：表格中第一个 text 类型字段视为去重 key。
    """
    _unit_re = re.compile(r'[（(]([^）)]{1,10})[）)]')
    fields = []
    seen_names = set()
    first_text_field = None

    for raw, name in zip(raw_fields, resolved_fields):
        if name in seen_names:
            logger.warning(f"字段名重复: '{raw}' 解析为 '{name}'（已存在），回退使用原始列名")
            name = raw
        seen_names.add(name)

        m = _unit_re.search(raw)
        unit = m.group(1).strip() if m else ""
        field_type = _infer_field_type(name, unit)

        f: dict = {"name": name, "type": field_type}
        if unit:
            f["unit"] = unit
        fields.append(f)

        if first_text_field is None and field_type == "text":
            first_text_field = name

    # 关键字段启发式：第一个 text 字段作为 dedup key 并标记 required
    dedup_key_fields = []
    if first_text_field:
        for f in fields:
            if f["name"] == first_text_field:
                f["required"] = True
                dedup_key_fields.append(first_text_field)
                break

    return fields, dedup_key_fields


def enrich_word_multi_table_specs(table_specs: list) -> None:
    """为每个 Word 表生成子 profile（table_profile）与内置抽取段落（builtin_prompt）。

    与表数量一一对应：几个表就有几个 table_profile 和几段 builtin_prompt。
    """
    for i, spec in enumerate(table_specs):
        if not isinstance(spec, dict):
            continue
        raw = spec.get("field_names") or []
        if not raw:
            continue
        resolved = resolve_field_names(raw)
        fields, dedup_key_fields = _enrich_fields(raw, resolved)
        idx = int(spec.get("table_index", i))
        above = (spec.get("instruction_above") or "").strip()
        instr = f"从源文档中提取可填入第 {idx + 1} 张 Word 表的数据（列名：{'、'.join(resolved[:32])}）"
        if above:
            instr += f"。表上方说明：{above[:500]}"
        tp: Dict[str, Any] = {
            "table_index": idx,
            "fields": fields,
            "instruction": instr,
        }
        if dedup_key_fields:
            tp["dedup_key_fields"] = dedup_key_fields
        spec["table_profile"] = tp

        type_hint = []
        for f in fields:
            d: Dict[str, Any] = {"name": f["name"], "type": f.get("type", "text")}
            if f.get("unit"):
                d["unit"] = f["unit"]
            type_hint.append(d)
        lines = [
            f"【第 {idx + 1} 张表 · 内置抽取提示】",
            f"表上方说明：{above}" if above else "表上方说明：（无）",
            f"本表列名（JSON 字段名须与下列一致，全局为各表并集）：{json.dumps(resolved, ensure_ascii=False)}",
            f"本表字段类型：{json.dumps(type_hint, ensure_ascii=False)}",
            "抽取时请结合本段说明理解该表语义；输出仍为全局 records，键为所有表列名并集，不适用的键填空字符串。",
        ]
        spec["builtin_prompt"] = "\n".join(lines)


_CN_NUM_MAP = {
    "一": 1, "二": 2, "三": 3, "四": 4, "五": 5,
    "六": 6, "七": 7, "八": 8, "九": 9, "十": 10,
}


def _table_no_to_index(token: str) -> Optional[int]:
    t = str(token or "").strip()
    if not t:
        return None
    if t.isdigit():
        v = int(t)
        return v - 1 if v > 0 else None
    if t in _CN_NUM_MAP:
        return _CN_NUM_MAP[t] - 1
    return None


def _extract_table_blocks(instruction: str) -> Dict[int, str]:
    text = str(instruction or "")
    if not text.strip():
        return {}
    blocks: Dict[int, str] = {}
    pattern = re.compile(
        r"(?:^|\n)\s*表([一二三四五六七八九十\d]+)\s*[：:]\s*(.*?)(?=(?:\n\s*表[一二三四五六七八九十\d]+\s*[：:])|\Z)",
        re.S,
    )
    for m in pattern.finditer(text):
        idx = _table_no_to_index(m.group(1))
        if idx is None:
            continue
        blocks[idx] = (m.group(2) or "").strip()
    return blocks


def apply_word_multi_instruction_constraints(profile: dict, instruction: str) -> dict:
    """把用户指令中的“分表约束”注入 table_specs（如城市/监测时间）。"""
    if not isinstance(profile, dict):
        return profile
    if profile.get("template_mode") != "word_multi_table":
        return profile

    specs = profile.get("table_specs")
    if not isinstance(specs, list) or not specs:
        return profile

    blocks = _extract_table_blocks(instruction or "")
    if not blocks:
        return profile

    city_re = re.compile(r"城市\s*[：:]\s*([^\s，,；;。]+)")
    time_re = re.compile(r"监测时间\s*[：:]\s*([^\n\r，,；;。]+)")
    out_specs = []
    for i, spec in enumerate(specs):
        if not isinstance(spec, dict):
            out_specs.append(spec)
            continue
        block = blocks.get(i, "")
        if not block:
            out_specs.append(spec)
            continue
        sp = dict(spec)
        fixed = dict(sp.get("fixed_values") or {})

        city_m = city_re.search(block)
        if city_m:
            city = city_m.group(1).strip()
            if city:
                sp["filter_field"] = "城市"
                sp["filter_value"] = city
                fixed["城市"] = city

        time_m = time_re.search(block)
        if time_m:
            mt = time_m.group(1).strip()
            if mt:
                fixed["监测时间"] = mt

        if fixed:
            sp["fixed_values"] = fixed
        out_specs.append(sp)

    out = dict(profile)
    out["table_specs"] = out_specs
    return out


def _extract_table_blocks(instruction: str) -> Dict[int, str]:
    text = str(instruction or "")
    if not text.strip():
        return {}
    blocks: Dict[int, str] = {}
    pattern = re.compile(
        r"(?:^|[\n\u3002\uff1b;])\s*\u8868\s*(["
        r"\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341\d]+)"
        r"\s*[:\uff1a]\s*(.*?)(?=(?:[\n\u3002\uff1b;]\s*\u8868\s*["
        r"\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341\d]+\s*[:\uff1a])|\Z)",
        re.S,
    )
    for m in pattern.finditer(text):
        idx = _table_no_to_index(m.group(1))
        if idx is None:
            continue
        blocks[idx] = (m.group(2) or "").strip()
    return blocks


def apply_word_multi_instruction_constraints(profile: dict, instruction: str) -> dict:
    """Inject per-table constraints like city/time into word_multi_table specs."""
    if not isinstance(profile, dict):
        return profile
    if profile.get("template_mode") != "word_multi_table":
        return profile

    specs = profile.get("table_specs")
    if not isinstance(specs, list) or not specs:
        return profile

    blocks = _extract_table_blocks(instruction or "")
    if not blocks:
        return profile

    city_field = "\u57ce\u5e02"
    time_field = "\u76d1\u6d4b\u65f6\u95f4"
    city_re = re.compile(
        r"\u57ce\u5e02\s*[:\uff1a=]\s*([^\n\r\uff1b;,\uff0c\u3002]+)"
    )
    time_re = re.compile(
        r"\u76d1\u6d4b\u65f6\u95f4\s*[:\uff1a=]\s*([^\n\r\uff1b;,\uff0c\u3002]+)"
    )

    out_specs = []
    for i, spec in enumerate(specs):
        if not isinstance(spec, dict):
            out_specs.append(spec)
            continue
        block = blocks.get(i, "")
        if not block:
            out_specs.append(spec)
            continue

        sp = dict(spec)
        fixed = dict(sp.get("fixed_values") or {})

        city_m = city_re.search(block)
        if city_m:
            city = city_m.group(1).strip()
            if city:
                sp["filter_field"] = city_field
                sp["filter_value"] = city
                fixed[city_field] = city

        time_m = time_re.search(block)
        if time_m:
            mt = time_m.group(1).strip()
            if mt:
                fixed[time_field] = mt

        if fixed:
            sp["fixed_values"] = fixed
        out_specs.append(sp)

    out = dict(profile)
    out["table_specs"] = out_specs
    return out


def apply_instruction_runtime_hints(profile: dict, instruction: str) -> dict:
    """
    根据用户指令为 profile 注入通用运行时提示（不改变用户原始语义）。
    当前仅做日期范围信号增强：若指令含日期区间且字段未覆盖日期，则补充日期字段。
    """
    if not isinstance(profile, dict):
        return profile

    instr = str(instruction or "").strip()
    out = dict(profile)
    if instr:
        out["instruction"] = instr

    task_mode = out.get("task_mode", "single_record")
    if task_mode != "table_records":
        return out
    if not parse_date_range_from_instruction(out.get("instruction", "")):
        return out

    fields = out.get("fields", [])
    if not isinstance(fields, list):
        fields = []
    names = {str(f.get("name", "")).strip() for f in fields if isinstance(f, dict)}
    if not any(n in names for n in ("日期", "统计日期", "监测时间", "时间", "date", "Date", "DATE")):
        out["fields"] = list(fields) + [{"name": "日期", "type": "text"}]
    return out


def generate_profile_from_template(
    template_path: str = None,
    use_llm: bool = False,
    mode: str = "auto",
    user_description: str = None,
) -> dict:
    """从模板文件或自然语言描述生成 profile

    优先级：template_path（文件解析） > user_description（LLM） > 默认字段
    """
    if mode in ("file", "auto") and template_path:
        try:
            detected = detect_template_structure(template_path)
            raw_fields = detected.pop("field_names", [])
            # 文件模板中的表头应优先按原始列名保留。
            # 否则别名/模糊归一化可能把明确模板列误改成其它业务字段，
            # 进而导致结构化抽取和写表全链路错位。
            resolved = list(raw_fields)

            fields, dedup_key_fields = _enrich_fields(raw_fields, resolved)

            instruction = _default_instruction(detected.get("task_mode", "table_records"), resolved)
            if detected.get("template_mode") == "word_multi_table" and detected.get("table_specs"):
                enrich_word_multi_table_specs(detected["table_specs"])
                instruction = _default_instruction_word_multi(resolved, detected["table_specs"])

            profile = {
                "report_name": Path(template_path).stem,
                "template_path": _rel_path(template_path),
                "instruction": instruction,
                "task_mode": detected.get("task_mode", "table_records"),
                "template_mode": detected.get("template_mode", "excel_table"),
                "fields": fields,
            }
            if dedup_key_fields:
                profile["dedup_key_fields"] = dedup_key_fields
            if detected.get("template_mode") == "word_multi_table" and detected.get("table_specs"):
                profile["table_profiles"] = [
                    s["table_profile"] for s in detected["table_specs"]
                    if isinstance(s, dict) and "table_profile" in s
                ]
                profile["word_multi_parallel"] = True
            profile.update({k: v for k, v in detected.items() if k not in profile})
            profile = apply_instruction_runtime_hints(profile, user_description or "")
            if profile.get("template_mode") == "word_multi_table":
                profile = apply_word_multi_instruction_constraints(profile, profile.get("instruction", ""))
            return profile
        except Exception as e:
            logger.warning(f"模板文件解析失败: {e}")
            if mode == "file":
                raise

    if user_description:
        return _profile_from_llm(user_description, template_path)

    return _default_profile(template_path)


def generate_profile_smart(
    template_path: str,
    instruction: str,
    document_sample: str = "",
) -> dict:
    """智能 profile 生成：LLM 分析模板内容 + 用户指令 + 文档样本"""
    template_content = _read_template_text(template_path) if template_path else "（无模板文件）"
    doc_part = f"\n\n【输入文档样本】\n{document_sample[:2000]}" if document_sample.strip() else ""

    prompt = f"""你是数据抽取配置专家。请分析以下模板和用户指令，生成JSON格式的抽取配置。

【模板内容】
{template_content}

【用户指令】
{instruction}{doc_part}

输出JSON，必须包含：
- "task_mode": "table_records" 或 "single_record"
- "template_mode": "excel_table" / "word_table" / "vertical"
- "fields": 字段数组，每个字段包含：
  - "name": 字段名（简洁明确）
  - "type": 字段类型，从以下选择：text（文本）、number（数值）、date（日期）、percentage（百分比/比率）、phone（电话号码）
  - "unit": 单位（可选，如"亿元"、"万人"、"%"）
- "instruction": 详细抽取说明

只输出JSON，不要其他文字："""

    try:
        from src.adapters.model_client import call_model
        result = call_model(prompt)
        if isinstance(result, dict) and result.get("fields"):
            return _normalize_profile(result, template_path, instruction)
    except Exception as e:
        logger.error(f"智能 profile 生成失败: {e}")

    # 回退到规则模式
    try:
        return generate_profile_from_template(template_path=template_path, mode="auto")
    except Exception:
        return _default_profile(template_path)


# ── 内部工具 ─────────────────────────────────────────────────────────────────

def _profile_from_llm(description: str, template_path: Optional[str]) -> dict:
    """使用 LLM 从自然语言描述生成 profile"""
    try:
        from src.adapters.model_client import call_model
        prompt = (
            f"请根据以下描述生成数据抽取字段列表（JSON格式）：\n{description}\n\n"
            f'输出格式：{{"fields": [{{"name": "字段名", "type": "类型", "unit": "单位（可选）"}}]}}\n'
            f'type 可选值：text（文本）、number（数值）、date（日期）、percentage（百分比/比率）、phone（电话号码）\n'
            f'只输出JSON，不要其他文字：'
        )
        result = call_model(prompt)
        if isinstance(result, dict) and result.get("fields"):
            fields = result["fields"]
        else:
            fields = [{"name": "字段1", "type": "text"}, {"name": "字段2", "type": "text"}]
    except Exception as e:
        logger.warning(f"LLM 字段生成失败: {e}")
        fields = [{"name": "字段1", "type": "text"}, {"name": "字段2", "type": "text"}]

    # 二次推断：LLM 返回 text 类型的字段尝试用规则推断为更精确的类型
    enriched_fields = []
    for f in fields:
        if not isinstance(f, dict) or "name" not in f:
            continue
        name = str(f["name"]).strip()
        if not name:
            continue
        unit = str(f.get("unit", "")).strip()
        declared_type = f.get("type", "text")
        if declared_type == "text":
            declared_type = _infer_field_type(name, unit)
        entry = {"name": name, "type": declared_type}
        if unit:
            entry["unit"] = unit
        enriched_fields.append(entry)

    return {
        "report_name": Path(template_path).stem if template_path else "llm_generated",
        "template_path": _rel_path(template_path) if template_path else "llm_generated",
        "instruction": f"从文档中提取：{description}",
        "task_mode": "single_record",
        "template_mode": "llm_generated",
        "fields": enriched_fields if enriched_fields else fields,
    }


def _default_profile(template_path: Optional[str]) -> dict:
    return {
        "report_name": Path(template_path).stem if template_path else "default",
        "template_path": _rel_path(template_path) if template_path else "default",
        "instruction": "从文档中提取关键信息",
        "task_mode": "single_record",
        "template_mode": "default",
        "fields": [
            {"name": "名称", "type": "text"},
            {"name": "数值", "type": "number"},
            {"name": "单位", "type": "text"},
            {"name": "备注", "type": "text"},
        ],
    }


def generate_profile_from_document(
    document_text: str,
    max_sample: int = 6000,
) -> dict:
    """从文档内容自动推断最佳表格结构（无模板、无用户指令时使用）

    内置 LLM prompt，全自动分析文档类型和字段结构：
    - 若文档包含重复结构（多条同类实体/记录）→ table_records 模式，自动推断字段
    - 若文档为单一主题（报告/合同/公文）→ single_record 模式，提取关键信息字段
    - 若文档过于杂乱无结构 → 返回 QA 模式标记，仅做信息入库

    Returns:
        profile dict，额外包含 "_doc_type" 字段：
        - "structured_table": 可制表（多条记录）
        - "structured_single": 可制表（单条记录）
        - "unstructured_qa": 杂乱文档，仅做 QA 入库
    """
    # 使用智能采样（15%：首部6% + 中部5% + 尾部4%）
    sample = _smart_sample_document(document_text) if document_text else ""
    if not sample.strip():
        return _default_profile(None)

    prompt = f"""分析以下文档样本，生成Profile和示例。

【文档样本】
{sample}

【任务】
1. 找出文档中每一个有具体数值的指标，每个指标作为一个字段
2. 为每个字段标注类型和单位（类型根据数值特征自动判断）
3. 根据这些字段，从样本中提取一个真实示例记录

【输出格式（严格JSON）】
{{
    "fields": [
        {{"name": "指标名称", "type": "类型", "unit": "单位"}}
    ],
    "example": {{
        "text": "从样本中选取的实际文本片段",
        "attributes": {{"指标名称": "实际值"}}
    }}
}}

只输出JSON，不要有其他内容。"""

    try:
        from src.adapters.model_client import call_model
        raw = call_model(prompt, temperature=0.1, timeout=30)

        # 兼容 call_model 返回的各种类型
        # 1. 如果是列表，取第一个元素（假设列表包含字典）
        if isinstance(raw, list):
            raw = raw[0] if raw else {}

        # 2. 如果是字符串，尝试提取 JSON 对象
        if isinstance(raw, str):
            import re
            match = re.search(r'\{.*\}', raw, re.DOTALL)
            if match:
                try:
                    raw = json.loads(match.group())
                except json.JSONDecodeError:
                    # 如果提取的 JSON 解析失败，尝试修复常见问题
                    try:
                        from src.adapters.model_client import _fix_json_common_issues
                        fixed = _fix_json_common_issues(match.group())
                        raw = json.loads(fixed)
                    except:
                        raw = {}
            else:
                # 没有找到 JSON 对象，设为空字典
                raw = {}

        # 3. 确保 raw 是字典，否则设为空字典
        if not isinstance(raw, dict):
            raw = {}

        # 新格式：提取fields和example
        fields_raw = raw.get("fields", [])
        example_data = raw.get("example", {})

        # 推断文档类型和任务模式
        # 默认假设为结构化单记录文档
        doc_type = "structured_single"
        task_mode = "single_record"
        instruction = "从文档中提取关键信息"

        # 简单启发式：如果字段数量多且包含常见表格字段，可能为表格
        if len(fields_raw) > 8:
            # 检查是否有典型表格字段
            table_keywords = {"序号", "编号", "名称", "城市", "地区", "项目", "产品"}
            field_names = [str(f.get("name", "")).lower() for f in fields_raw if isinstance(f, dict)]
            if any(keyword in " ".join(field_names) for keyword in table_keywords):
                doc_type = "structured_table"
                task_mode = "table_records"
                instruction = "从文档中逐条提取记录"

        # QA 模式：返回最小化 profile + 标记
        if doc_type == "unstructured_qa":
            return {
                "report_name": "auto_qa",
                "template_path": "",
                "instruction": instruction or "将文档内容存储为知识库，供智能问答使用",
                "task_mode": "single_record",
                "template_mode": "generic",
                "_doc_type": "unstructured_qa",
                "fields": [
                    {"name": "文档标题", "type": "text"},
                    {"name": "摘要", "type": "text"},
                    {"name": "关键词", "type": "text"},
                    {"name": "主要内容", "type": "text"},
                ],
            }

        # 结构化模式：规范化字段
        fields = []
        seen = set()
        first_text_field = None
        for f in fields_raw:
            if not isinstance(f, dict) or "name" not in f:
                continue
            name = str(f["name"]).strip()
            if not name or name in seen:
                continue
            seen.add(name)
            unit = str(f["unit"]).strip() if f.get("unit") else ""
            # LLM 已给出类型，但若为 text 则尝试推断
            declared_type = f.get("type", "text")
            if declared_type == "text":
                declared_type = _infer_field_type(name, unit)
            entry = {"name": name, "type": declared_type}
            if unit:
                entry["unit"] = unit
            fields.append(entry)
            if first_text_field is None and declared_type == "text":
                first_text_field = name

        # 关键字段启发式：第一个 text 字段
        dedup_key_fields = []
        if first_text_field:
            for f in fields:
                if f["name"] == first_text_field:
                    f["required"] = True
                    dedup_key_fields.append(first_text_field)
                    break

        if not fields:
            raise ValueError("LLM 未生成有效字段")

        # 保存example数据到profile中
        example_attrs = {}
        example_text = ""
        if example_data and isinstance(example_data, dict):
            example_attrs = example_data.get("attributes", {})
            example_text = example_data.get("text", "")
            # 确保attributes是字典
            if not isinstance(example_attrs, dict):
                example_attrs = {}
            # 如果example_text为空但attributes有值，生成简单文本
            if not example_text and example_attrs:
                example_text = "示例记录: " + ", ".join([f"{k}: {v}" for k, v in example_attrs.items()])

        profile = {
            "report_name": "auto_generated",
            "template_path": "",
            "instruction": instruction,
            "task_mode": task_mode if doc_type == "structured_table" else "single_record",
            "template_mode": "generic",
            "_doc_type": doc_type,
            "fields": fields,
            "_example": example_attrs,
            "_example_text": example_text,
        }
        if dedup_key_fields:
            profile["dedup_key_fields"] = dedup_key_fields

        estimated = raw.get("estimated_record_count")
        if estimated and isinstance(estimated, (int, float)):
            profile["_estimated_record_count"] = int(estimated)

        logger.info(f"自动文档分析: type={doc_type}, fields={len(fields)}, task_mode={profile['task_mode']}")
        return profile

    except Exception as e:
        logger.warning(f"文档自动分析失败: {e}，使用默认 profile")
        # 回退：使用规则生成函数
        return _rule_generate_profile_and_example(sample)


def _heuristic_profile_from_text(text: str) -> dict:
    """简单启发式文档分析（LLM 不可用时的回退）"""
    lines = text.strip().split('\n')
    # 检测是否有表格式结构（多行相似分隔符）
    separator_lines = sum(1 for l in lines if '|' in l or '\t' in l)
    numbered_lines = sum(1 for l in lines if re.match(r'^\s*\d+[\.\、]', l))

    if separator_lines > 3 or numbered_lines > 5:
        # 疑似表格或列表结构
        return {
            "report_name": "auto_generated",
            "template_path": "",
            "instruction": "从文档中提取所有条目的关键信息，每个条目一条记录",
            "task_mode": "table_records",
            "template_mode": "generic",
            "_doc_type": "structured_table",
            "fields": [
                {"name": "序号", "type": "text"},
                {"name": "名称", "type": "text"},
                {"name": "描述", "type": "text"},
                {"name": "数值", "type": "number"},
                {"name": "备注", "type": "text"},
            ],
        }

    return {
        "report_name": "auto_generated",
        "template_path": "",
        "instruction": "从文档中提取关键信息",
        "task_mode": "single_record",
        "template_mode": "generic",
        "_doc_type": "structured_single",
        "fields": [
            {"name": "标题", "type": "text"},
            {"name": "主题", "type": "text"},
            {"name": "关键信息", "type": "text"},
            {"name": "数据摘要", "type": "text"},
            {"name": "备注", "type": "text"},
        ],
    }


def _default_instruction(task_mode: str, fields) -> str:
    names = "、".join(fields[:6]) if fields else "各字段"
    if task_mode == "table_records":
        return f"从文档表格中逐行提取记录，包含字段：{names}"
    return f"从文档中提取以下信息：{names}"


def _default_instruction_word_multi(resolved: list, table_specs: list) -> str:
    """多表 Word：结合各表上方说明与每表列名生成默认抽取指令。"""
    lines = [
        "从源文档中抽取表格数据；模板含多个 Word 表格，请同时遵守「各表上方说明」与「各表列名（表头）」。"
    ]
    for i, spec in enumerate(table_specs):
        if not isinstance(spec, dict):
            continue
        idx = int(spec.get("table_index", i))
        above = (spec.get("instruction_above") or "").strip()
        cols = spec.get("field_names") or []
        col_s = "、".join(cols[:24]) if cols else ""
        if above:
            lines.append(f"表格{idx + 1} — 表上方说明：{above[:800]}")
        if col_s:
            lines.append(f"表格{idx + 1} — 该表应填列名：{col_s}")
    tail = "、".join(resolved[:24]) if resolved else "各字段"
    lines.append(
        f"输出 records 的字段名为以上各表列名的并集，包含：{tail}；仅适用于某张表的列在其它表对应的记录中可留空。"
    )
    return "\n".join(lines)


def _rel_path(path: str) -> str:
    try:
        rel = os.path.relpath(path)
        return rel if ".." not in rel and len(rel) <= 100 else Path(path).name
    except Exception:
        return Path(path).name


def _normalize_profile(result: dict, template_path: Optional[str], instruction: str) -> dict:
    """标准化 LLM 返回的 profile，补充类型推断和关键字段识别"""
    fields = []
    seen = set()
    first_text_field = None
    _unit_re = re.compile(r'[（(]([^）)]{1,10})[）)]')
    for f in result.get("fields", []):
        if not isinstance(f, dict) or "name" not in f:
            continue
        name = str(f["name"]).strip()
        if not name or name in seen:
            continue
        seen.add(name)
        # 保留 LLM 推断的类型，若为 text 则尝试自动推断
        unit = f.get("unit", "")
        if not unit:
            m = _unit_re.search(name)
            if m:
                unit = m.group(1).strip()
        declared_type = f.get("type", "text")
        if declared_type == "text":
            declared_type = _infer_field_type(name, unit)
        entry = {"name": name, "type": declared_type}
        if unit:
            entry["unit"] = unit
        if f.get("required"):
            entry["required"] = True
        fields.append(entry)
        if first_text_field is None and declared_type == "text":
            first_text_field = name

    # 关键字段启发式：第一个 text 字段
    dedup_key_fields = []
    if first_text_field:
        for f in fields:
            if f["name"] == first_text_field:
                f["required"] = True
                dedup_key_fields.append(first_text_field)
                break

    profile = {
        "report_name": Path(template_path).stem if template_path else "auto",
        "template_path": _rel_path(template_path) if template_path else "auto",
        "instruction": result.get("instruction") or instruction,
        "task_mode": result.get("task_mode", "table_records"),
        "template_mode": result.get("template_mode", "excel_table"),
        "fields": fields,
        "header_row": result.get("header_row", 1),
        "start_row": result.get("start_row", 2),
    }
    if dedup_key_fields:
        profile["dedup_key_fields"] = dedup_key_fields
    return profile


def _read_template_text(template_path: str) -> str:
    ext = Path(template_path).suffix.lower()
    if ext in (".xlsx", ".xlsm"):
        try:
            from openpyxl import load_workbook
            wb = load_workbook(template_path, read_only=True, data_only=True)
            ws = wb.active
            lines = []
            for r in range(1, min((ws.max_row or 30), 30) + 1):
                row = [str(ws.cell(r, c).value or "").strip()
                       for c in range(1, min((ws.max_column or 20), 20) + 1)]
                row = [v for v in row if v]
                if row:
                    lines.append(" | ".join(row))
            return "\n".join(lines)
        except Exception as e:
            return f"（Excel 解析失败: {e}）"
    if ext == ".docx":
        try:
            from docx import Document
            doc = Document(template_path)
            parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()][:20]
            for i, t in enumerate(doc.tables[:3]):
                parts.append(f"【表格{i+1}】")
                for row in t.rows[:3]:
                    parts.append(" | ".join(c.text.strip() for c in row.cells))
            return "\n".join(parts)
        except Exception as e:
            return f"（Word 解析失败: {e}）"
    return f"（不支持格式: {ext}）"


# ── 智能采样函数 ──────────────────────────────────────────────────────────────

def _smart_sample_document(document_text: str, total_ratio: float = 0.15) -> str:
    """智能采样文档内容，总计15%：首部6% + 中部5% + 尾部4%

    Args:
        document_text: 完整文档文本
        total_ratio: 总采样比例，默认0.15（15%）

    Returns:
        采样后的文本
    """
    if len(document_text) < 5000:
        return document_text  # 短文档使用全文

    # 1. 按段落分割
    paragraphs = re.split(r'\n\s*\n|\r\n\s*\r\n', document_text)
    if len(paragraphs) <= 1:
        paragraphs = document_text.split('\n')

    # 2. 计算目标字符数
    target_chars = int(len(document_text) * total_ratio)
    head_ratio, middle_ratio, tail_ratio = 0.06, 0.05, 0.04

    # 3. 采样策略
    sampled_paragraphs = []

    # 首部6%：从开头取段落
    head_chars = int(len(document_text) * head_ratio)
    head_paras = _collect_paragraphs_up_to_chars(paragraphs, head_chars)
    sampled_paragraphs.extend(head_paras)

    # 尾部4%：从结尾取段落
    tail_chars = int(len(document_text) * tail_ratio)
    tail_paras = _collect_paragraphs_from_end_up_to_chars(paragraphs, tail_chars)
    sampled_paragraphs.extend(tail_paras)

    # 中部5%：优先选择包含数字的段落
    middle_chars = int(len(document_text) * middle_ratio)
    middle_indices = range(len(paragraphs) // 4, 3 * len(paragraphs) // 4)
    middle_paras = []

    # 优先选择包含数字的段落
    numeric_paras = []
    for idx in middle_indices:
        if idx < len(paragraphs) and re.search(r'\d+', paragraphs[idx]):
            numeric_paras.append(paragraphs[idx])

    # 如果包含数字的段落不足，补充其他段落
    if len(numeric_paras) < 3:  # 至少3段
        other_paras = []
        for idx in middle_indices:
            if idx < len(paragraphs):
                para = paragraphs[idx]
                if para not in numeric_paras:
                    other_paras.append(para)
        numeric_paras.extend(other_paras[:5])  # 最多补充5段

    middle_paras = _collect_paragraphs_up_to_chars(numeric_paras, middle_chars)
    sampled_paragraphs.extend(middle_paras)

    # 4. 去重并保持顺序
    unique_paras = []
    seen = set()
    for para in sampled_paragraphs:
        if para not in seen and para.strip():
            unique_paras.append(para)
            seen.add(para)

    # 5. 回退：如果采样结果为空，取前15%
    if not unique_paras:
        fallback_chars = int(len(document_text) * 0.15)
        return document_text[:fallback_chars]

    return "\n\n".join(unique_paras)


def _collect_paragraphs_up_to_chars(paragraphs: list, max_chars: int) -> list:
    """按顺序收集段落，直到达到目标字符数

    Args:
        paragraphs: 段落列表
        max_chars: 最大字符数

    Returns:
        收集的段落列表
    """
    result = []
    current_chars = 0
    for para in paragraphs:
        if current_chars + len(para) <= max_chars:
            result.append(para)
            current_chars += len(para)
        else:
            break
    return result


def _collect_paragraphs_from_end_up_to_chars(paragraphs: list, max_chars: int) -> list:
    """从末尾收集段落，直到达到目标字符数

    Args:
        paragraphs: 段落列表
        max_chars: 最大字符数

    Returns:
        收集的段落列表（保持原文顺序）
    """
    result = []
    current_chars = 0
    for para in reversed(paragraphs):
        if current_chars + len(para) <= max_chars:
            result.insert(0, para)  # 保持原文顺序
            current_chars += len(para)
        else:
            break
    return result


def _rule_generate_profile_and_example(sample_text: str) -> dict:
    """LLM失败时的规则回退函数

    Args:
        sample_text: 采样文本

    Returns:
        包含fields和example的profile字典
    """
    # 正则提取 数字+单位
    pattern = r'(\d[\d,.]*)\s*([元%人万吨公里小时]*)'
    matches = re.findall(pattern, sample_text)

    fields = []
    example_attrs = {}

    for i, (num, unit) in enumerate(matches[:10]):  # 最多10个字段
        field_name = f"指标{i+1}"
        field_type = "text"

        if unit:
            if unit == "元":
                field_type = "money"
            elif unit == "%":
                field_type = "percentage"
            elif unit in ("人", "万人"):
                field_type = "population"
            elif unit in ("吨", "万吨"):
                field_type = "weight"
            else:
                field_type = "number"

        fields.append({
            "name": field_name,
            "type": field_type,
            "unit": unit
        })
        example_attrs[field_name] = num + unit

    # 构造示例文本
    example_text = sample_text[:500] if len(sample_text) > 500 else sample_text

    # 构建完整profile
    doc_type = "structured_single"
    task_mode = "single_record"
    instruction = "从文档中提取关键信息"

    # 简单启发式：如果字段数量多，可能为表格
    if len(fields) > 5:
        doc_type = "structured_table"
        task_mode = "table_records"
        instruction = "从文档中逐条提取记录"

    # 关键字段启发式：第一个字段
    dedup_key_fields = []
    if fields and fields[0]["type"] == "text":
        dedup_key_fields.append(fields[0]["name"])
        fields[0]["required"] = True

    profile = {
        "report_name": "rule_generated",
        "template_path": "",
        "instruction": instruction,
        "task_mode": task_mode,
        "template_mode": "generic",
        "_doc_type": doc_type,
        "fields": fields,
        "_example": example_attrs,
        "_example_text": example_text,
    }

    if dedup_key_fields:
        profile["dedup_key_fields"] = dedup_key_fields

    return profile
