"""
文档智能操作模块（M1）

将自然语言指令解析为结构化操作，并在 Excel / Word 文档上执行。

支持的操作类型：
  format_cells   — 设置字体（加粗/颜色/大小）、对齐、边框、背景色
  edit_content   — 替换文本内容、删除行/列、插入内容
  adjust_layout  — 调整列宽/行高、合并/拆分单元格
  filter_rows    — 按条件筛选/删除行
  extract_data   — 按条件提取数据（不修改文档，返回记录）
"""

import json
import logging
import re
import shutil
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────
# 指令解析
# ─────────────────────────────────────────────

_PARSE_PROMPT = """\
你是一个文档操作助手。将用户的自然语言指令解析为结构化 JSON 操作命令。

文档信息：
{doc_info}

用户指令：{instruction}

请输出一个 JSON 对象，格式如下（只输出 JSON，不要解释）：
{{
  "operation": "format_cells|edit_content|adjust_layout|filter_rows|extract_data",
  "target": {{
    "sheet": "工作表名（Excel 专用，默认 null）",
    "range": "单元格范围如 A1:D10（null 表示整表）",
    "column": "列名称或字母如 C 或 城市（null 表示不限列）",
    "row": "行号（整数，null 表示不限行）",
    "keyword": "目标关键词（用于定位文本）"
  }},
  "action": "bold|italic|color|bg_color|font_size|align|border|replace|delete_row|delete_col|merge|col_width|row_height|filter_keep|filter_delete",
  "params": {{
    "bold": true,
    "italic": false,
    "color": "#FF0000（hex 颜色，null 表示不改）",
    "bg_color": "#FFFF00",
    "font_size": 12,
    "align": "left|center|right|top|middle|bottom",
    "value": "替换后的内容（edit_content 时使用）",
    "condition_field": "条件字段名（filter_rows 时使用）",
    "condition_op": ">|<|>=|<=|==|!=|contains|startswith|endswith",
    "condition_value": "条件值",
    "width": 20,
    "height": 20
  }}
}}

只输出 JSON："""


def _build_doc_info(document_path: str) -> str:
    """生成文档的简要描述，供 LLM 理解文档结构"""
    p = Path(document_path)
    ext = p.suffix.lower()
    info_lines = [f"文件名：{p.name}", f"格式：{ext}"]

    if ext in ('.xlsx', '.xls', '.xlsm'):
        try:
            from openpyxl import load_workbook
            wb = load_workbook(document_path, read_only=True, data_only=True)
            for ws in wb.worksheets[:3]:
                headers = []
                for cell in next(ws.iter_rows(max_row=1), []):
                    if cell.value is not None:
                        headers.append(str(cell.value))
                info_lines.append(f"工作表 {ws.title}：列标题={headers[:10]}, 行数≈{ws.max_row}")
            wb.close()
        except Exception as e:
            info_lines.append(f"Excel 信息获取失败: {e}")

    elif ext in ('.docx', '.doc'):
        try:
            from docx import Document
            doc = Document(document_path)
            info_lines.append(f"段落数：{len(doc.paragraphs)}, 表格数：{len(doc.tables)}")
            if doc.tables:
                t = doc.tables[0]
                if t.rows:
                    headers = [c.text.strip() for c in t.rows[0].cells[:8]]
                    info_lines.append(f"第一个表格列标题：{headers}")
        except Exception as e:
            info_lines.append(f"Word 信息获取失败: {e}")

    return "\n".join(info_lines)


def parse_instruction(instruction: str, document_path: str) -> Dict[str, Any]:
    """将自然语言指令解析为结构化操作命令

    Returns:
        {"operation": ..., "target": {...}, "action": ..., "params": {...}}
    """
    doc_info = _build_doc_info(document_path)
    prompt = _PARSE_PROMPT.format(doc_info=doc_info, instruction=instruction)

    try:
        from src.adapters.model_client import call_model
        result = call_model(prompt)
        if isinstance(result, dict):
            return result
        # 尝试从字符串中提取 JSON
        text = str(result)
        m = re.search(r'\{[\s\S]*\}', text)
        if m:
            return json.loads(m.group())
    except Exception as e:
        logger.warning(f"指令解析失败，使用启发式解析: {e}")

    return _heuristic_parse(instruction)


def _heuristic_parse(instruction: str) -> Dict[str, Any]:
    """无 LLM 时的启发式指令解析"""
    cmd: Dict[str, Any] = {
        "operation": "format_cells",
        "target": {"sheet": None, "range": None, "column": None, "row": None, "keyword": None},
        "action": "bold",
        "params": {}
    }

    instr = instruction.lower()

    # 操作类型识别
    if any(k in instr for k in ('加粗', 'bold', '粗体')):
        cmd["action"] = "bold"
        cmd["params"]["bold"] = True
    elif any(k in instr for k in ('颜色', 'color', '字体颜色', '背景')):
        cmd["action"] = "color"
        if '红' in instr: cmd["params"]["color"] = "#FF0000"
        elif '蓝' in instr: cmd["params"]["color"] = "#0000FF"
        elif '绿' in instr: cmd["params"]["color"] = "#00AA00"
        elif '黄' in instr: cmd["params"]["bg_color"] = "#FFFF00"
    elif any(k in instr for k in ('居中', 'center', '对齐')):
        cmd["action"] = "align"
        if '左' in instr: cmd["params"]["align"] = "left"
        elif '右' in instr: cmd["params"]["align"] = "right"
        else: cmd["params"]["align"] = "center"
    elif any(k in instr for k in ('列宽', '宽度', 'width')):
        cmd["operation"] = "adjust_layout"
        cmd["action"] = "col_width"
        nums = re.findall(r'\d+', instr)
        cmd["params"]["width"] = int(nums[0]) if nums else 20
    elif any(k in instr for k in ('删除', 'delete', '移除')):
        cmd["operation"] = "edit_content"
        cmd["action"] = "delete_row"
    elif any(k in instr for k in ('替换', 'replace', '修改为', '改为')):
        cmd["operation"] = "edit_content"
        cmd["action"] = "replace"
    elif any(k in instr for k in ('筛选', 'filter', '过滤', '保留', '删除.*行', '行.*删除')):
        cmd["operation"] = "filter_rows"
        cmd["action"] = "filter_delete"
    elif any(k in instr for k in ('提取', 'extract', '找出', '查找')):
        cmd["operation"] = "extract_data"
        cmd["action"] = "filter_keep"

    # 列识别
    col_m = re.search(r'第([一二三四五六七八九十\d]+)[列柱]|([A-Z])[列柱]', instruction)
    if col_m:
        cn = col_m.group(1) or col_m.group(2)
        cn_map = {'一': 'A', '二': 'B', '三': 'C', '四': 'D', '五': 'E',
                  '六': 'F', '七': 'G', '八': 'H', '九': 'I', '十': 'J'}
        cmd["target"]["column"] = cn_map.get(cn, cn)

    # 行识别
    row_m = re.search(r'第(\d+)[行]', instruction)
    if row_m:
        cmd["target"]["row"] = int(row_m.group(1))

    return cmd


# ─────────────────────────────────────────────
# Excel 操作执行器
# ─────────────────────────────────────────────

def _col_letter_to_index(col: Optional[str], ws) -> Optional[int]:
    """将列字母/名称转为 1-based 列索引，支持通过列标题名查找"""
    if col is None:
        return None
    # 字母列标
    if re.match(r'^[A-Za-z]+$', col):
        from openpyxl.utils import column_index_from_string
        try:
            return column_index_from_string(col.upper())
        except Exception:
            pass
    # 按第一行列标题名查找
    first_row = list(ws.iter_rows(min_row=1, max_row=1, values_only=True))
    if first_row:
        for i, val in enumerate(first_row[0], start=1):
            if val is not None and str(val).strip() == str(col).strip():
                return i
    # 纯数字
    if str(col).isdigit():
        return int(col)
    return None


def _apply_cell_format(cell, params: Dict[str, Any]):
    """将格式参数应用到单个单元格"""
    from openpyxl.styles import Font, Alignment, PatternFill
    from openpyxl.styles.colors import Color

    font_kwargs = {}
    if "bold" in params:
        font_kwargs["bold"] = params["bold"]
    if "italic" in params:
        font_kwargs["italic"] = params["italic"]
    if "font_size" in params and params["font_size"]:
        font_kwargs["size"] = params["font_size"]
    if "color" in params and params["color"]:
        try:
            font_kwargs["color"] = params["color"].lstrip('#')
        except Exception:
            pass

    if font_kwargs:
        existing = cell.font
        cell.font = Font(
            name=existing.name,
            bold=font_kwargs.get("bold", existing.bold),
            italic=font_kwargs.get("italic", existing.italic),
            size=font_kwargs.get("size", existing.size),
            color=font_kwargs.get("color", None) or (existing.color.rgb if existing.color and existing.color.type == 'rgb' else None),
        )

    if "align" in params and params["align"]:
        h_map = {"left": "left", "center": "center", "right": "right"}
        v_map = {"top": "top", "middle": "center", "bottom": "bottom"}
        h = h_map.get(params["align"], params["align"])
        v = v_map.get(params["align"])
        cell.alignment = Alignment(
            horizontal=h if h in ("left", "center", "right") else None,
            vertical=v if v in ("top", "center", "bottom") else None,
        )

    if "bg_color" in params and params["bg_color"]:
        color = params["bg_color"].lstrip('#')
        cell.fill = PatternFill(fill_type="solid", fgColor=color)


def _parse_range(ws, target: Dict[str, Any]) -> Tuple[int, int, int, int]:
    """解析 target 为 (min_row, max_row, min_col, max_col)"""
    if target.get("range"):
        from openpyxl.utils import range_boundaries
        try:
            min_col, min_row, max_col, max_row = range_boundaries(target["range"].upper())
            return min_row, max_row, min_col, max_col
        except Exception:
            pass
    col_idx = _col_letter_to_index(target.get("column"), ws)
    row_idx = target.get("row")
    min_row = row_idx if row_idx else 2
    max_row = ws.max_row
    min_col = col_idx if col_idx else 1
    max_col = col_idx if col_idx else ws.max_column
    return min_row, max_row, min_col, max_col


def _eval_condition(cell_value, op: str, cond_value: Any) -> bool:
    """计算单元格值是否满足条件"""
    try:
        v = cell_value
        cv = cond_value

        # 数值比较
        if op in ('>', '<', '>=', '<='):
            v = float(str(v).replace(',', '').strip())
            cv = float(str(cv).replace(',', '').strip())
            if op == '>':  return v > cv
            if op == '<':  return v < cv
            if op == '>=': return v >= cv
            if op == '<=': return v <= cv

        # 字符串比较
        sv, scv = str(v).strip(), str(cv).strip()
        if op == '==':         return sv == scv
        if op == '!=':         return sv != scv
        if op == 'contains':   return scv in sv
        if op == 'startswith': return sv.startswith(scv)
        if op == 'endswith':   return sv.endswith(scv)
    except Exception:
        pass
    return False


def _execute_excel(document_path: str, output_path: str, command: Dict[str, Any]) -> Dict[str, Any]:
    """在 Excel 文档上执行操作命令"""
    from openpyxl import load_workbook
    wb = load_workbook(document_path)

    target = command.get("target", {})
    action = command.get("action", "")
    params = command.get("params", {})
    operation = command.get("operation", "")

    # 选取工作表
    sheet_name = target.get("sheet")
    ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active

    affected = 0

    # ── 格式操作 ──
    if operation == "format_cells" or action in (
        "bold", "italic", "color", "bg_color", "font_size", "align", "border"
    ):
        min_row, max_row, min_col, max_col = _parse_range(ws, target)
        for row in ws.iter_rows(min_row=min_row, max_row=max_row, min_col=min_col, max_col=max_col):
            for cell in row:
                _apply_cell_format(cell, params)
                affected += 1

    # ── 替换内容 ──
    elif action == "replace":
        keyword = target.get("keyword")
        new_val = params.get("value", "")
        for row in ws.iter_rows():
            for cell in row:
                if cell.value is not None:
                    if keyword and keyword in str(cell.value):
                        cell.value = str(cell.value).replace(keyword, new_val)
                        affected += 1
                    elif not keyword:
                        cell.value = new_val
                        affected += 1

    # ── 列宽调整 ──
    elif action == "col_width":
        col_idx = _col_letter_to_index(target.get("column"), ws)
        width = params.get("width", 15)
        if col_idx:
            from openpyxl.utils import get_column_letter
            ws.column_dimensions[get_column_letter(col_idx)].width = width
            affected = 1
        else:
            for col_letter in ws.column_dimensions:
                ws.column_dimensions[col_letter].width = width
                affected += 1

    # ── 行高调整 ──
    elif action == "row_height":
        min_row, max_row, _, _ = _parse_range(ws, target)
        height = params.get("height", 20)
        for r in range(min_row, max_row + 1):
            ws.row_dimensions[r].height = height
            affected += 1

    # ── 行删除 ──
    elif action == "delete_row":
        row_idx = target.get("row")
        if row_idx:
            ws.delete_rows(row_idx)
            affected = 1

    # ── 列删除 ──
    elif action == "delete_col":
        col_idx = _col_letter_to_index(target.get("column"), ws)
        if col_idx:
            ws.delete_cols(col_idx)
            affected = 1

    # ── 行筛选：按条件保留/删除 ──
    elif action in ("filter_keep", "filter_delete") or operation == "filter_rows":
        cond_field = params.get("condition_field")
        cond_op = params.get("condition_op", "==")
        cond_val = params.get("condition_value", "")

        # 找到条件字段列索引
        header_row = list(ws.iter_rows(min_row=1, max_row=1, values_only=True))
        col_idx = None
        if header_row and cond_field:
            for i, h in enumerate(header_row[0], start=1):
                if h is not None and str(h).strip() == str(cond_field).strip():
                    col_idx = i
                    break

        if col_idx is None and target.get("column"):
            col_idx = _col_letter_to_index(target["column"], ws)

        if col_idx:
            keep_action = (action == "filter_keep")
            rows_to_delete = []
            for row in ws.iter_rows(min_row=2):
                cell_val = row[col_idx - 1].value
                match = _eval_condition(cell_val, cond_op, cond_val)
                if keep_action and not match:
                    rows_to_delete.append(row[0].row)
                elif not keep_action and match:
                    rows_to_delete.append(row[0].row)
            for r in reversed(rows_to_delete):
                ws.delete_rows(r)
                affected += 1

    # ── 合并单元格 ──
    elif action == "merge":
        rng = target.get("range")
        if rng:
            ws.merge_cells(rng)
            affected = 1

    # ── 提取数据（不修改文档）──
    elif operation == "extract_data":
        cond_field = params.get("condition_field")
        cond_op = params.get("condition_op", "==")
        cond_val = params.get("condition_value", "")
        header_row = list(ws.iter_rows(min_row=1, max_row=1, values_only=True))
        headers = [str(h) if h is not None else "" for h in (header_row[0] if header_row else [])]

        col_idx = None
        if cond_field and headers:
            for i, h in enumerate(headers):
                if h.strip() == str(cond_field).strip():
                    col_idx = i
                    break

        records = []
        for row in ws.iter_rows(min_row=2, values_only=True):
            if col_idx is not None:
                if not _eval_condition(row[col_idx], cond_op, cond_val):
                    continue
            rec = {headers[i]: row[i] for i in range(min(len(headers), len(row)))}
            records.append(rec)
            affected += 1

        wb.close()
        return {"status": "ok", "operation": "extract_data", "records": records, "count": len(records)}

    wb.save(output_path)
    wb.close()
    logger.info(f"Excel 操作完成: {action}, 影响 {affected} 个元素 → {output_path}")
    return {"status": "ok", "operation": action, "affected": affected, "output_path": output_path}


# ─────────────────────────────────────────────
# Word 操作执行器
# ─────────────────────────────────────────────

def _execute_word(document_path: str, output_path: str, command: Dict[str, Any]) -> Dict[str, Any]:
    """在 Word 文档上执行操作命令"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document(document_path)
    target = command.get("target", {})
    action = command.get("action", "")
    params = command.get("params", {})
    operation = command.get("operation", "")

    keyword = target.get("keyword") or params.get("value", "")
    affected = 0

    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }

    # ── 格式操作（段落 / 跑文） ──
    if operation == "format_cells" or action in ("bold", "italic", "font_size", "color", "align"):
        for para in doc.paragraphs:
            if keyword and keyword not in para.text:
                continue
            if "align" in params and params["align"]:
                para.alignment = align_map.get(params["align"], WD_ALIGN_PARAGRAPH.LEFT)
                affected += 1
            for run in para.runs:
                if "bold" in params:
                    run.bold = params["bold"]
                if "italic" in params:
                    run.italic = params["italic"]
                if "font_size" in params and params["font_size"]:
                    run.font.size = Pt(params["font_size"])
                if "color" in params and params["color"]:
                    hex_c = params["color"].lstrip('#')
                    r, g, b = int(hex_c[0:2], 16), int(hex_c[2:4], 16), int(hex_c[4:6], 16)
                    run.font.color.rgb = RGBColor(r, g, b)
                affected += 1

    # ── 文本替换 ──
    elif action == "replace":
        old_text = target.get("keyword", "")
        new_text = params.get("value", "")
        for para in doc.paragraphs:
            if old_text in para.text:
                for run in para.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
                        affected += 1
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
                                affected += 1

    # ── 删除段落（包含关键词的段落）──
    elif action == "delete_row":
        keyword_to_del = target.get("keyword", "")
        to_delete = [p for p in doc.paragraphs if keyword_to_del and keyword_to_del in p.text]
        for p in to_delete:
            p._element.getparent().remove(p._element)
            affected += 1

    # ── Word 表格操作 ──
    elif operation in ("filter_rows", "extract_data") and doc.tables:
        table = doc.tables[0]
        cond_field = params.get("condition_field")
        cond_op = params.get("condition_op", "==")
        cond_val = params.get("condition_value", "")

        headers = [cell.text.strip() for cell in table.rows[0].cells]
        col_idx = None
        if cond_field:
            for i, h in enumerate(headers):
                if h == cond_field:
                    col_idx = i
                    break

        if operation == "extract_data":
            records = []
            for row in table.rows[1:]:
                vals = [c.text.strip() for c in row.cells]
                if col_idx is not None:
                    if not _eval_condition(vals[col_idx], cond_op, cond_val):
                        continue
                records.append(dict(zip(headers, vals)))
                affected += 1
            return {"status": "ok", "operation": "extract_data", "records": records, "count": len(records)}

        elif operation == "filter_rows" and col_idx is not None:
            rows_to_del = []
            for row in table.rows[1:]:
                vals = [c.text.strip() for c in row.cells]
                match = _eval_condition(vals[col_idx], cond_op, cond_val)
                if action == "filter_keep" and not match:
                    rows_to_del.append(row)
                elif action == "filter_delete" and match:
                    rows_to_del.append(row)
            for row in rows_to_del:
                row._element.getparent().remove(row._element)
                affected += 1

    doc.save(output_path)
    logger.info(f"Word 操作完成: {action}, 影响 {affected} 个元素 → {output_path}")
    return {"status": "ok", "operation": action, "affected": affected, "output_path": output_path}


# ─────────────────────────────────────────────
# 主入口
# ─────────────────────────────────────────────

def operate_document(
    instruction: str,
    document_path: str,
    output_path: Optional[str] = None,
    backup: bool = True,
) -> Dict[str, Any]:
    """将自然语言指令转化为文档操作并执行

    Args:
        instruction:    自然语言指令，如"将第三列加粗"
        document_path:  源文档路径
        output_path:    输出文档路径（None 则覆盖原文件）
        backup:         执行前备份原文件

    Returns:
        {
          "status": "ok" | "error",
          "output_path": str,
          "operation": str,
          "affected": int,
          "records": [...],   # extract_data 时返回
          "backup_path": str, # backup=True 时返回
          "command": {...},   # 解析后的操作命令
        }
    """
    doc_path = Path(document_path)
    if not doc_path.exists():
        return {"status": "error", "message": f"文件不存在: {document_path}"}

    ext = doc_path.suffix.lower()
    if ext not in ('.xlsx', '.xls', '.xlsm', '.docx', '.doc'):
        return {"status": "error", "message": f"不支持的文件格式: {ext}，仅支持 Excel/Word"}

    # 确定输出路径
    if output_path is None:
        output_path = str(doc_path)
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)

    # 备份原文件
    backup_path = None
    if backup:
        backup_path = str(doc_path.parent / f"{doc_path.stem}_backup{doc_path.suffix}")
        shutil.copy2(document_path, backup_path)

    # 解析指令
    try:
        command = parse_instruction(instruction, document_path)
        logger.info(f"指令解析结果: {command}")
    except Exception as e:
        return {"status": "error", "message": f"指令解析失败: {e}"}

    # 执行操作
    try:
        if ext in ('.xlsx', '.xls', '.xlsm'):
            result = _execute_excel(document_path, output_path, command)
        else:
            result = _execute_word(document_path, output_path, command)

        result["backup_path"] = backup_path
        result["command"] = command
        result["output_path"] = output_path
        return result

    except Exception as e:
        logger.error(f"文档操作执行失败: {e}", exc_info=True)
        # 回滚
        if backup_path and Path(backup_path).exists():
            shutil.copy2(backup_path, document_path)
        return {"status": "error", "message": str(e), "command": command}


def rollback(backup_path: str, target_path: str) -> bool:
    """将备份文件恢复到目标路径"""
    try:
        shutil.copy2(backup_path, target_path)
        logger.info(f"回滚完成: {backup_path} → {target_path}")
        return True
    except Exception as e:
        logger.error(f"回滚失败: {e}")
        return False
