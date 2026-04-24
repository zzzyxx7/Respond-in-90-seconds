from pathlib import Path
from typing import List

from openpyxl import load_workbook
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph


def _iter_body_blocks(doc: Document):
    """按文档顺序遍历正文中的段落与表格（与 doc.tables 顺序一致）。"""
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield DocxTable(child, doc)


def _paragraphs_text_before_each_table(doc: Document) -> List[str]:
    """每个表格之前累积的段落文本（表与表之间的说明文字）。"""
    current: List[str] = []
    before: List[str] = []
    for block in _iter_body_blocks(doc):
        if isinstance(block, Paragraph):
            t = (block.text or "").strip()
            if t:
                current.append(t)
        else:
            before.append("\n".join(current))
            current = []
    return before


def _header_row_fields(table: DocxTable) -> List[str]:
    if len(table.rows) < 1:
        return []
    header_fields = []
    for cell in table.rows[0].cells:
        text = (cell.text or "").strip()
        if text:
            header_fields.append(text)
    return header_fields


def _merge_unique_field_names(per_table: List[List[str]]) -> List[str]:
    seen = set()
    out: List[str] = []
    for row in per_table:
        for name in row:
            if name not in seen:
                seen.add(name)
                out.append(name)
    return out


def detect_template_structure(template_path: str, multi_table: bool = False) -> dict:
    """检测模板结构

    Args:
        template_path: 模板文件路径
        multi_table: 是否启用多表格检测（仅对Word文档有效）

    Returns:
        dict: 模板结构信息
    """
    ext = Path(template_path).suffix.lower()
    if ext in [".xlsx", ".xlsm"]:
        return detect_excel_structure(template_path)
    if ext == ".docx":
        return detect_word_structure(template_path, multi_table=multi_table)
    raise ValueError(f"暂不支持的模板类型：{ext}")


def detect_excel_structure(template_path: str) -> dict:
    wb = load_workbook(template_path)
    ws = wb.active
    max_row = ws.max_row
    max_col = ws.max_column

    a_values, b_values = [], []
    for r in range(1, min(max_row, 50) + 1):
        a_values.append(ws.cell(r, 1).value)
        b_values.append(ws.cell(r, 2).value if max_col >= 2 else None)

    a_non_empty = [str(v).strip() for v in a_values if v is not None and str(v).strip()]
    b_empty_ratio = sum(1 for v in b_values if v is None or not str(v).strip()) / max(1, len(b_values))

    if len(a_non_empty) >= 3 and b_empty_ratio >= 0.5:
        field_names = []
        start_row = 1
        first = str(ws.cell(1,1).value or '').strip()
        second = str(ws.cell(1,2).value or '').strip() if max_col >= 2 else ''
        if first in {'字段', 'Field', '指标'} and second in {'值', 'Value', '内容', ''}:
            start_row = 2
        for r in range(start_row, min(max_row, 100) + 1):
            v = ws.cell(r, 1).value
            if v is not None and str(v).strip():
                field_names.append(str(v).strip())
        return {'task_mode': 'single_record', 'template_mode': 'vertical', 'field_names': field_names}

    for r in range(1, min(max_row, 15) + 1):
        row_values = [ws.cell(r, c).value for c in range(1, max_col + 1)]
        row_texts = [str(v).strip() for v in row_values if v is not None and str(v).strip()]
        if len(row_texts) < 2:
            continue
        short_count = sum(1 for x in row_texts if len(x) <= 20)
        if short_count >= 2:
            return {'task_mode': 'table_records', 'template_mode': 'excel_table', 'header_row': r, 'start_row': r + 1, 'field_names': row_texts}
    raise ValueError('无法自动识别 Excel 模板结构')


def detect_word_structure(template_path: str, multi_table: bool = False) -> dict:
    """检测 Word 模板结构。

    - 单个表格：template_mode=word_table，取首行表头。
    - 多个表格：template_mode=word_multi_table，为每个表记录首行表头及表上方连续段落文字（填表说明）。

    Args:
        multi_table: 保留参数（与 Excel 多表等 API 对齐）；当前以 ``len(doc.tables) > 1`` 自动启用多表检测。
    """
    _ = multi_table
    doc = Document(template_path)
    if not doc.tables:
        raise ValueError("Word 模板中没有表格，暂不支持自动识别纯正文模板")
    if len(doc.tables) > 1:
        return _detect_multi_word_table(template_path)
    return _detect_single_word_table(template_path)


def _detect_multi_word_table(template_path: str) -> dict:
    """多表 Word：每表第一行为列名；表与上一表之间的段落合并为 instruction_above。"""
    doc = Document(template_path)
    if not doc.tables:
        raise ValueError("Word 模板中没有表格")
    before_texts = _paragraphs_text_before_each_table(doc)
    if len(before_texts) != len(doc.tables):
        # 与 doc.tables 一一对应；异常时回退为等长空串
        while len(before_texts) < len(doc.tables):
            before_texts.append("")
        before_texts = before_texts[: len(doc.tables)]

    table_specs = []
    per_table_fields: List[List[str]] = []
    for i, table in enumerate(doc.tables):
        fields_i = _header_row_fields(table)
        if len(fields_i) < 1:
            raise ValueError(f"第 {i + 1} 个 Word 表格无法识别表头（首行无有效列名）")
        per_table_fields.append(fields_i)
        above = (before_texts[i] or "").strip()
        spec = {
            "table_index": i,
            "field_names": fields_i,
            "instruction_above": above,
            # 模板预留数据行容量（不含表头），用于控制填表超量。
            "max_rows": max(0, len(table.rows) - 1),
        }
        if above:
            first_line = above.split("\n", 1)[0].strip()
            if first_line:
                spec["description"] = first_line[:120]
        table_specs.append(spec)

    union_fields = _merge_unique_field_names(per_table_fields)
    if len(union_fields) < 1:
        raise ValueError("多表 Word 模板未识别到任何列名")

    return {
        "task_mode": "table_records",
        "template_mode": "word_multi_table",
        "table_index": 0,
        "header_row": 0,
        "start_row": 1,
        "field_names": union_fields,
        "table_specs": table_specs,
    }


def _detect_single_word_table(template_path: str) -> dict:
    """检测Word文档中的单个表格（原有逻辑）"""
    doc = Document(template_path)
    if not doc.tables:
        raise ValueError('Word 模板中没有表格，暂不支持自动识别纯正文模板')
    table = doc.tables[0]
    if len(table.rows) < 1:
        raise ValueError('Word 模板表格为空')
    header_fields = []
    for cell in table.rows[0].cells:
        text = cell.text.strip()
        if text:
            header_fields.append(text)
    if len(header_fields) < 2:
        raise ValueError('无法识别 Word 表格表头')
    return {'task_mode': 'table_records', 'template_mode': 'word_table', 'table_index': 0, 'header_row': 0, 'start_row': 1, 'field_names': header_fields}
