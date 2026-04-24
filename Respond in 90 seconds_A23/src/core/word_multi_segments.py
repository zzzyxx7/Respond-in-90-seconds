"""
多表 Word 方案 B：为每张模板表构造一段源文档上下文（优先从输入 .docx 按表切分）。
"""

from __future__ import annotations

import logging
from typing import Any, Dict, List, Optional

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)


def _iter_body_blocks(doc: Document):
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield DocxTable(child, doc)


def _paragraphs_before_each_table(doc: Document) -> List[str]:
    current: List[str] = []
    before: List[str] = []
    for block in _iter_body_blocks(doc):
        if isinstance(block, Paragraph):
            t = (block.text or "").strip()
            if t:
                current.append(t)
        else:
            before.append("\n".join(current))
            current = []
    return before


def _table_cells_text(table: DocxTable) -> str:
    lines: List[str] = []
    for row in table.rows:
        cells = [(c.text or "").strip() for c in row.cells]
        if any(cells):
            lines.append("\t".join(cells))
    return "\n".join(lines)


def segments_from_input_docx(docx_path: str, n_tables: int) -> Optional[List[str]]:
    """若 docx 中表格数 >= n_tables，为前 n_tables 张表各生成一段上下文：表上说明 + 表内文本。"""
    try:
        doc = Document(docx_path)
    except Exception as e:
        logger.warning("无法打开输入 docx %s: %s", docx_path, e)
        return None
    if len(doc.tables) < n_tables:
        return None
    before_texts = _paragraphs_before_each_table(doc)
    while len(before_texts) < len(doc.tables):
        before_texts.append("")
    before_texts = before_texts[: len(doc.tables)]

    segments: List[str] = []
    for i in range(n_tables):
        above = (before_texts[i] or "").strip()
        body = _table_cells_text(doc.tables[i])
        parts = []
        if above:
            parts.append(f"【表{i + 1}上方说明】\n{above}")
        if body:
            parts.append(f"【表{i + 1}内容】\n{body}")
        segments.append("\n\n".join(parts) if parts else "")
    return segments


def _first_docx_path(documents: List[Dict[str, Any]]) -> Optional[str]:
    for doc in documents:
        if not isinstance(doc, dict):
            continue
        p = doc.get("path") or doc.get("file_path") or ""
        if p and str(p).lower().endswith(".docx"):
            return str(p)
    return None


def build_word_multi_table_segments(
    profile: Dict[str, Any],
    all_text: str,
    documents: List[Dict[str, Any]],
) -> List[str]:
    """返回与 table_specs 等长的文本片段列表；无法按表切分时每段均为全文。"""
    specs = profile.get("table_specs") or []
    n = len(specs)
    if n == 0:
        return [all_text]

    path = _first_docx_path(documents or [])
    if path:
        segs = segments_from_input_docx(path, n)
        if segs and len(segs) == n:
            # 若某段过短（无表格正文），用全文补足，避免空上下文
            out = []
            for i, seg in enumerate(segs):
                s = (seg or "").strip()
                if len(s) < 80 and all_text.strip():
                    out.append(all_text)
                else:
                    out.append(seg if seg else all_text)
            return out

    logger.info(
        "word_multi_segments: 未找到可匹配的输入 docx 或表数量不一致，每表使用全文作为上下文（仍并行单表 prompt）"
    )
    return [all_text] * n
