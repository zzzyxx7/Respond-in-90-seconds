from copy import copy
import logging
from openpyxl import load_workbook
from docx import Document

_logger = logging.getLogger(__name__)


def _normalize_records(records):
    payload = records if isinstance(records, dict) else None
    if isinstance(records, dict) and 'records' in records:
        records = records['records']

    if not isinstance(records, list):
        raise ValueError("表格填充要求数据为 list[dict] 或 {'records': [...]} 格式")

    for item in records:
        if not isinstance(item, dict):
            raise ValueError('records 中每一项都必须是 dict')

    return records, payload


def create_excel_from_records(output_path: str, records: list):
    """动态创建Excel文件，从记录字段自动推断列名（用于无模板场景）"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill

    # 规范化records
    if isinstance(records, dict) and 'records' in records:
        records = records['records']
    if not isinstance(records, list):
        records = []

    # 收集所有字段名（保持顺序）
    all_fields = []
    seen = set()
    for record in records:
        if isinstance(record, dict):
            for key in record.keys():
                if key not in seen and not str(key).startswith('_'):
                    all_fields.append(key)
                    seen.add(key)

    wb = Workbook()
    ws = wb.active

    if not all_fields:
        wb.save(output_path)
        return

    # 写表头（加粗）
    header_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
    for col, field in enumerate(all_fields, 1):
        cell = ws.cell(row=1, column=col, value=field)
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal='center')
        cell.fill = header_fill

    # 写数据
    for row_idx, record in enumerate(records, 2):
        if not isinstance(record, dict):
            continue
        for col, field in enumerate(all_fields, 1):
            value = record.get(field, '')
            ws.cell(row=row_idx, column=col, value=value)

    # 自动调整列宽
    for col in range(1, len(all_fields) + 1):
        max_len = 0
        for row in ws.iter_rows(min_row=1, max_row=min(ws.max_row, 20), min_col=col, max_col=col):
            for cell in row:
                if cell.value:
                    max_len = max(max_len, len(str(cell.value)))
        ws.column_dimensions[ws.cell(row=1, column=col).column_letter].width = min(max_len + 4, 40)

    wb.save(output_path)


def fill_excel_vertical(template_path: str, output_path: str, data: dict):
    # 加载字段别名映射（用于模板字段名与规范化字段名之间的映射）
    reverse_alias_map = {}
    try:
        from src.core.alias import load_alias_map, build_reverse_alias_map
        alias_map = load_alias_map()
        reverse_alias_map = build_reverse_alias_map(alias_map)
    except Exception as e:
        import logging
        logging.getLogger(__name__).warning("字段别名映射加载失败: %s，将继续使用原始字段名匹配", e)

    wb = load_workbook(template_path)
    ws = wb.active

    for row in range(2, ws.max_row + 1):
        field_name = ws[f'A{row}'].value
        if field_name is None:
            continue
        field_name = str(field_name).strip()

        # 1. 尝试直接获取
        value = data.get(field_name, '')

        # 2. 如果直接获取失败，尝试通过字段别名映射查找
        if not value and field_name in reverse_alias_map:
            normalized_name = reverse_alias_map[field_name]
            value = data.get(normalized_name, '')

        # 3. 如果规范化名称查找失败，尝试模糊匹配（大小写不敏感）
        if not value:
            for data_key in data.keys():
                if data_key.lower() == field_name.lower():
                    value = data.get(data_key, '')
                    break

        if value:
            ws[f'B{row}'] = value

    wb.save(output_path)


def fill_excel_table(template_path: str, output_path: str, records, header_row: int = 1, start_row: int = 2):
    # 加载字段别名映射（用于模板字段名与规范化字段名之间的映射）
    reverse_alias_map = {}
    try:
        from src.core.alias import load_alias_map, build_reverse_alias_map
        alias_map = load_alias_map()
        reverse_alias_map = build_reverse_alias_map(alias_map)
    except Exception as e:
        import logging
        logging.getLogger(__name__).warning("字段别名映射加载失败: %s，将继续使用原始字段名匹配", e)

    records, _ = _normalize_records(records)

    if not records:
        _logger.warning(
            "excel_table 填充：抽取记录为空，输出将仅保留表头（请检查 LLM 抽取结果与模板表头字段是否匹配）。output=%s",
            output_path,
        )

    wb = load_workbook(template_path)
    ws = wb.active

    header_map = {}
    for col in range(1, ws.max_column + 1):
        cell_value = ws.cell(row=header_row, column=col).value
        if cell_value is not None:
            header_map[str(cell_value).strip()] = col

    if not header_map:
        raise ValueError('Excel 模板表头为空，无法进行表格填充')

    # 优化：批量插入行并复制样式，然后批量写入数据
    # 计算需要的总行数
    total_rows_needed = start_row + len(records) - 1
    rows_to_insert = max(0, total_rows_needed - ws.max_row)

    if rows_to_insert > 0:
        # 一次性插入所有需要的行
        ws.insert_rows(start_row + 1, amount=rows_to_insert)

        # 复制样式：为新插入的行复制模板行的样式
        template_style_row = start_row
        for row_offset in range(rows_to_insert):
            target_row = start_row + 1 + row_offset
            for col in range(1, ws.max_column + 1):
                src = ws.cell(row=template_style_row, column=col)
                dst = ws.cell(row=target_row, column=col)
                if src.has_style:
                    dst._style = copy(src._style)
                dst.font = copy(src.font)
                dst.fill = copy(src.fill)
                dst.border = copy(src.border)
                dst.alignment = copy(src.alignment)
                dst.protection = copy(src.protection)
                dst.number_format = src.number_format

    # 批量写入数据（支持字段别名映射）
    import re
    _unit_suffix_re = re.compile(r'[（(][^）)]*[）)]')

    for i, record in enumerate(records):
        target_row = start_row + i
        for field_name, col in header_map.items():
            # 1. 尝试直接获取
            value = record.get(field_name, '')

            # 2. 如果直接获取失败，尝试通过字段别名映射查找
            if not value and field_name in reverse_alias_map:
                normalized_name = reverse_alias_map[field_name]
                value = record.get(normalized_name, '')

            # 3. 去掉括号单位后再查别名（如 "常住人口（万）" → "常住人口" → 人口）
            if not value:
                stripped = _unit_suffix_re.sub('', field_name).strip()
                if stripped != field_name:
                    value = record.get(stripped, '')
                    if not value and stripped in reverse_alias_map:
                        normalized_name = reverse_alias_map[stripped]
                        value = record.get(normalized_name, '')

            # 4. 反向查找：用 record 的每个 key 查别名，看是否对应同一个规范名
            if not value:
                field_canonical = reverse_alias_map.get(field_name) or reverse_alias_map.get(
                    _unit_suffix_re.sub('', field_name).strip())
                if field_canonical:
                    for record_key in record.keys():
                        record_canonical = reverse_alias_map.get(record_key)
                        if record_canonical and record_canonical == field_canonical:
                            value = record.get(record_key, '')
                            break

            # 5. 大小写不敏感匹配（兜底）
            if not value:
                for record_key in record.keys():
                    if record_key.lower() == field_name.lower():
                        value = record.get(record_key, '')
                        break

            ws.cell(row=target_row, column=col, value=value)

    wb.save(output_path)


def fill_word_table(template_path: str, output_path: str, records, table_index: int = 0, header_row: int = 0, start_row: int = 1):
    records, payload = _normalize_records(records)

    doc = Document(template_path)

    if not doc.tables:
        raise ValueError('Word 模板中没有找到表格')

    table_groups = payload.get('_table_groups') if isinstance(payload, dict) else None
    if isinstance(table_groups, list) and table_groups:
        for group in table_groups:
            idx = int(group.get('table_index', 0))
            if idx >= len(doc.tables):
                continue
            group_records = group.get('records', [])
            # 无记录时不写入、不清空数据行，避免把整张表刷成空白
            if not group_records:
                continue
            # 如果records是字符串，尝试解析为JSON
            if isinstance(group_records, str):
                try:
                    import json
                    group_records = json.loads(group_records)
                except:
                    # 如果解析失败，跳过这个表格
                    import logging
                    logging.getLogger(__name__).warning(
                        "表%s的records是字符串且无法解析为JSON: %s",
                        idx,
                        str(group_records)[:100],
                    )
                    continue
            _fill_single_word_table(doc.tables[idx], group_records, header_row=header_row, start_row=start_row)
        doc.save(output_path)
        return

    if table_index >= len(doc.tables):
        raise ValueError(f'Word 模板只有 {len(doc.tables)} 个表格，table_index={table_index} 越界')

    _fill_single_word_table(doc.tables[table_index], records, header_row=header_row, start_row=start_row)
    doc.save(output_path)


def _fill_single_word_table(table, records, header_row: int = 0, start_row: int = 1):
    # 加载字段别名映射（用于模板字段名与规范化字段名之间的映射）
    reverse_alias_map = {}
    try:
        from src.core.alias import load_alias_map, build_reverse_alias_map
        alias_map = load_alias_map()
        reverse_alias_map = build_reverse_alias_map(alias_map)
    except Exception as e:
        import logging
        logging.getLogger(__name__).warning("字段别名映射加载失败: %s，将继续使用原始字段名匹配", e)

    header_map = {}
    header_cells = table.rows[header_row].cells
    for idx, cell in enumerate(header_cells):
        text = cell.text.strip()
        if text:
            header_map[text] = idx

    if not header_map:
        raise ValueError('Word 模板表头为空，无法进行表格填充')

    required_rows = start_row + len(records)
    while len(table.rows) < required_rows:
        table.add_row()

    for row_idx in range(start_row, len(table.rows)):
        row_cells = table.rows[row_idx].cells
        for cell in row_cells:
            cell.text = ''

    for i, record in enumerate(records):
        row_idx = start_row + i
        row_cells = table.rows[row_idx].cells
        # 确保record是字典
        if not isinstance(record, dict):
            import logging
            logging.getLogger(__name__).warning(
                "记录%s不是字典类型: %s, 值: %s",
                i,
                type(record),
                str(record)[:100],
            )
            continue

        # 遍历表头字段，查找对应的记录值（支持字段别名映射）
        for field_name, col_idx in header_map.items():
            # 1. 尝试直接获取
            value = record.get(field_name, '')

            # 2. 如果直接获取失败，尝试通过字段别名映射查找
            if not value and field_name in reverse_alias_map:
                normalized_name = reverse_alias_map[field_name]
                value = record.get(normalized_name, '')

            # 3. 如果规范化名称查找失败，尝试模糊匹配（大小写不敏感）
            if not value:
                for record_key in record.keys():
                    if record_key.lower() == field_name.lower():
                        value = record.get(record_key, '')
                        break

            if value is not None:
                row_cells[col_idx].text = str(value)


